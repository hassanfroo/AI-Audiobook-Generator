from docx import Document

doc = Document()
doc.add_heading('Test Academic Paper', 0)

doc.add_heading('Chapter 1: Introduction', level=1)
doc.add_paragraph('This is the first chapter of our test. It should be extracted as a separate file in chapter mode [1].')

doc.add_heading('Chapter 2: Methods', level=1)
doc.add_paragraph('This is the second chapter. (Smith, 2026). It will prove that docx and chapter splitting work.')

doc.save('test_input.docx')
print("Created test_input.docx")
